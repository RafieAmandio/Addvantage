#!/usr/bin/env python3
"""Generate a professionally styled ANTS / DOMAIN Product Document as .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# Brand colors — dark navy + amber accent for "alpha intelligence" / tactical feel
PRIMARY = RGBColor(0x0A, 0x0F, 0x1E)      # Near-black navy
ACCENT = RGBColor(0xF5, 0x9E, 0x0B)        # Amber
ACCENT_DARK = RGBColor(0xB4, 0x53, 0x09)   # Dark amber / burnt orange
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF8, 0xFA, 0xFC)
MED_GRAY = RGBColor(0x64, 0x74, 0x8B)
DARK_TEXT = RGBColor(0x1E, 0x29, 0x3B)
BORDER_COLOR = "F59E0B"
HEADER_FILL = "0A0F1E"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_TEXT
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
        + ''.join(f'<w:{k} w:val="single" w:sz="4" w:space="0" w:color="{v}"/>' for k, v in kwargs.items())
        + '</w:tcBorders>')
    tcPr.append(tcBorders)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY if level == 1 else ACCENT_DARK
        run.font.name = 'Calibri'
    if level == 1:
        # Accent line under H1
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{BORDER_COLOR}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    return h


def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, HEADER_FILL)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_TEXT
            bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
            set_cell_shading(cell, bg)

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT
    else:
        p.clear()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT
    return p


def add_para(text, italic=False, size=11, after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_TEXT
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.4
    return p


def add_info_box(text, accent=False):
    """Colored callout box. accent=True uses amber theme."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = ACCENT_DARK if accent else PRIMARY
    run.font.bold = True
    fill = "FEF3C7" if accent else "F1F5F9"
    border = BORDER_COLOR if accent else "CBD5E1"
    set_cell_shading(cell, fill)
    set_cell_border(cell, top=border, bottom=border, left=border, right=border)
    doc.add_paragraph()


def add_quote(text):
    """Pull-quote style block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="{BORDER_COLOR}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = PRIMARY


# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════

for _ in range(5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

# ANTS title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ANTS")
run.font.size = Pt(64)
run.font.bold = True
run.font.color.rgb = PRIMARY
run.font.name = 'Calibri'

# ANTS expansion
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Alpha Network Trading System")
run.font.size = Pt(16)
run.font.color.rgb = MED_GRAY
run.font.name = 'Calibri'

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(18)
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:bottom w:val="single" w:sz="18" w:space="1" w:color="{BORDER_COLOR}"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# DOMAIN
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DOMAIN")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = ACCENT_DARK
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Directional Outlook & Macro Alpha Intelligence")
run.font.size = Pt(13)
run.font.italic = True
run.font.color.rgb = MED_GRAY
run.font.name = 'Calibri'

# Tagline
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Welcome to the DOMAIN — Alpha Intelligence for global markets."')
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = DARK_TEXT
run.font.name = 'Calibri'

# Doc type
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Product Specification Document")
run.font.size = Pt(14)
run.font.color.rgb = MED_GRAY
run.font.name = 'Calibri'

# Version block
for _ in range(2):
    doc.add_paragraph()

info_lines = [
    ("Version", "2.0"),
    ("Date", "April 2026"),
    ("Status", "Draft"),
    ("Classification", "Confidential"),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}: ")
    run.font.size = Pt(11)
    run.font.color.rgb = MED_GRAY
    run.font.name = 'Calibri'
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = DARK_TEXT
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════

add_heading_styled("Executive Summary", level=1)

add_para(
    "ANTS — Alpha Network Trading System — is a one-way market intelligence platform built for "
    "experienced traders. Its product surface, DOMAIN (Directional Outlook & Macro Alpha Intelligence), "
    "delivers a continuous stream of unfiltered news, real-time market impact analysis, an economic "
    "calendar, a daily/weekly trading plan, structured trading education, and 1-on-1 consultation with "
    "AI and the team. There is no community, no chat, and no public forum — by design.",
    after=10,
)
add_para(
    "Distribution is Telegram-first: free pillars (news, calendar, the founder's personal channel, "
    "psychology primers) live on Telegram, and Telegram is the funnel that pushes users into the paid "
    "DOMAIN web app, where the trading plan, 1v1 consultation, and full education library are gated. "
    "The product is monetised through a freemium model with tiered packages billed in 3-month "
    "increments (e.g. VIP+ Trader), deliberately filtering out churners and tourists.",
    after=10,
)
add_para(
    "ANTS is explicit about who it is not for — beginners, gamblers, and hobbyists are turned away at "
    "the door. The audience is experienced market participants who want a copilot to sharpen their edge "
    "and activate \"Six Eyes\" on the market. Every paid user signs a liability agreement at the end of "
    "the upgrade flow that releases ANTS from responsibility for trading losses. A built-in trading "
    "psychology layer — named theory primers (Kahneman's Loss Aversion, prospect theory, recency bias) "
    "and a hashtag system for common trader pathologies (#losing-streak, #unprofitability, "
    "#developing-edge) — connects news, the trading plan, education, and consultation into a single "
    "searchable surface.",
    after=10,
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════

add_heading_styled("Table of Contents", level=1)

toc_items = [
    "Executive Summary",
    "1. Product Overview",
    "2. Positioning",
    "3. Information Pillars",
    "4. Trading Psychology Layer",
    "5. Freemium & Tier Model",
    "6. Distribution: Telegram-First Funnel",
    "7. Influencer & Collab Channels",
    "8. Anti-Community Philosophy",
    "9. Liability Agreement",
    "10. Key Surfaces & Screens",
    "11. User Roles",
    "12. Open Questions",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT_DARK
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. PRODUCT OVERVIEW
# ═══════════════════════════════════════════════════════════

add_heading_styled("1. Product Overview", level=1)

add_para(
    "ANTS is the brand. DOMAIN is the product surface that paid users interact with. "
    "The entire system is built on a one-way information model: the team broadcasts, the user receives. "
    "No community, no chat, no comments — by design, not by omission."
)

add_styled_table(
    ["Attribute", "Detail"],
    [
        ["Brand", "ANTS — Alpha Network Trading System"],
        ["Product Surface", "DOMAIN — Directional Outlook & Macro Alpha Intelligence"],
        ["Concept", "Market radar + copilot powered by AI and professional traders"],
        ["Information Model", "One-way broadcast — no community, chat, or comments"],
        ["Primary Funnel", "Telegram (everything routes through Telegram)"],
        ["Secondary Surface", "Web app (DOMAIN dashboard) for paid tiers"],
        ["Monetisation", "Freemium with tiered paid packages billed in 3-month increments"],
        ["Audience", "Experienced traders and serious market participants only"],
    ],
    col_widths=[2.0, 4.5],
)

add_quote(
    "ANTS is not a research subscription, a signals group, or a trading school. It is a market "
    "radar — a one-way intelligence stream meant to sharpen the edge of people who already trade."
)

# ═══════════════════════════════════════════════════════════
# 2. POSITIONING
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled("2. Positioning", level=1)

add_heading_styled("What is this?", level=2)
add_quote(
    "A market radar powered by AI and professionals, created for traders and investors who already "
    "know what they're doing."
)

add_heading_styled("Who is this NOT for?", level=2)
add_bullet("Beginners")
add_bullet("Gamblers")
add_bullet("Hobbyists")
add_para(
    "We do not soften this. The product is built around the assumption that the user already "
    "understands risk, position sizing, and market structure. Onboarding should filter, not convert.",
    italic=True,
)

add_heading_styled("Who is this ACTUALLY for?", level=2)
add_bullet("Experienced market participants trying to level up their edge")
add_bullet("Traders who need a market copilot for second-opinion / blind-spot coverage")
add_bullet('Anyone who wants to activate "Six Eyes" on the market — continuous, multi-angle situational awareness')

# ═══════════════════════════════════════════════════════════
# 3. INFORMATION PILLARS
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled("3. Information Pillars", level=1)

add_para(
    "ANTS delivers six core streams. All are one-way. None of them have a comments section."
)

add_styled_table(
    ["#", "Pillar", "Description", "Tier"],
    [
        ["1", "Unfiltered News + Impact Analysis", "Raw market-moving news in real time, each item annotated with what it means for price", "Free"],
        ["2", "Economic Calendar", "Curated calendar of high-impact events, central bank decisions, earnings, macro releases", "Free"],
        ["3", "Trading Plan", "Daily / weekly directional trading plan from the team — setups, levels, invalidation", "LOCKED"],
        ["4", "AI + Team Consultation (1v1)", "Private 1-on-1 chat with the AI and the team for trade reviews and second opinions", "LOCKED"],
        ["5", "Education", "Structured trading education — process, psychology, risk, system design", "Free / Paid"],
        ["6", "My Channel", "Personal broadcast channel for the founder / lead voice — daily takes and market context", "Free"],
    ],
    col_widths=[0.4, 1.7, 3.5, 0.9],
)

add_heading_styled("Streaming (optional surface)", level=2)
add_bullet("If technically cheap: live stream from inside the DOMAIN web app")
add_bullet("Otherwise: stream from Telegram, which doubles as a conversion funnel back into the paid product")
add_bullet("Streaming is treated as a marketing surface, not a core deliverable")

# ═══════════════════════════════════════════════════════════
# 4. TRADING PSYCHOLOGY LAYER
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled("4. Trading Psychology Layer", level=1)

add_para(
    "A first-class part of the product, not a blog category. The goal is to give traders shared "
    "language for the failure modes they all hit."
)

add_heading_styled("Named Theory Primers", level=2)
add_para(
    "Short explainers built around real frameworks — Kahneman's Loss Aversion, prospect theory, "
    "recency bias, sunk-cost trap, and others. Each primer is tagged into the hashtag system so it "
    "can be pulled from inside news posts, the trading plan, and consultation logs."
)

add_heading_styled("Hashtag System", level=2)
add_para(
    "Every recurring trader pathology gets a tag that runs across posts, education, and consultation:"
)
add_bullet("losing-streak — process for surviving and exiting drawdown", bold_prefix="#")
add_bullet("unprofitability — diagnosing why an edge isn't there yet", bold_prefix="#")
add_bullet("developing-edge — building, testing, and trusting a new system", bold_prefix="#")
add_bullet("(more to be added as patterns emerge)")

add_info_box(
    "The hashtag system is the connective tissue between News, Trading Plan, Education, and "
    "Consultation. A user reading about a losing streak in education can pull every consultation "
    "log, post, and primer with the same tag.",
    accent=True,
)

# ═══════════════════════════════════════════════════════════
# 5. FREEMIUM & TIER MODEL
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled("5. Freemium & Tier Model", level=1)

add_heading_styled("Free Tier", level=2)
add_bullet("Unfiltered news + impact analysis")
add_bullet("Economic calendar")
add_bullet("My Channel (founder's broadcast)")
add_bullet("Some education content")
add_bullet("Public-facing psychology primers")

add_heading_styled("Paid Tier (Locked Features)", level=2)
add_bullet("Trading Plan — the daily/weekly directional plan", bold_prefix="")
add_bullet("AI + Team Consultation (1v1) — private chat", bold_prefix="")
add_bullet("Full education library")
add_bullet("Any future closed influencer collab channels")

add_heading_styled("Pricing Structure", level=2)
add_bullet("Billed in 3-month increments, not monthly — a deliberate filter against churners and tourists")
add_bullet('Tiered packages (e.g. "VIP+ Trader") with price increments per tier')
add_bullet("Exact tier names, prices, and per-tier gating: TBD (see Open Questions)")

add_info_box(
    "The 3-month billing window is a positioning choice. Anyone unwilling to commit a quarter is "
    "not the target audience.",
    accent=True,
)

# ═══════════════════════════════════════════════════════════
# 6. DISTRIBUTION: TELEGRAM-FIRST FUNNEL
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled("6. Distribution: Telegram-First Funnel", level=1)

add_para(
    "Telegram is THE funnel. Every other surface — web, streams, social — exists to push users "
    "into Telegram, and Telegram pushes them into the paid web product."
)

add_bullet("Telegram delivers all free pillars (news, calendar, my channel, psychology tags)")
add_bullet("Telegram hosts streams (or mirrors them) for conversion")
add_bullet("Telegram is where the call-to-action to upgrade lives")
add_bullet("The web app (DOMAIN dashboard) is the destination for paid users only — Trading Plan, 1v1 Consultation, full Education")

add_info_box(
    "Implication: the web product can be smaller and sharper than a typical SaaS. It does not need "
    "to handle discovery, onboarding noise, or community. It exists to host what Telegram cannot — "
    "gated content and private 1v1 surfaces.",
)

# ═══════════════════════════════════════════════════════════
# 7. INFLUENCER & COLLAB CHANNELS
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled("7. Influencer & Collab Channels", level=1)

add_para(
    "ANTS is designed to host guest voices without losing the one-way broadcast model."
)

add_bullet("An external influencer or partner trader can be given a dedicated channel inside the system to share their info / trades / takes")
add_bullet("Each collab channel is its own broadcast surface — still one-way, still no chat, still tagged into the same hashtag system")
add_bullet("This lets ANTS scale voice without scaling moderation, and gives partners a clean place to publish")

# ═══════════════════════════════════════════════════════════
# 8. ANTI-COMMUNITY PHILOSOPHY
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled("8. Anti-Community Philosophy", level=1)

add_para(
    "We are explicit and proud about NOT having a community / chat / forum. This is a feature, "
    "not an omission. The FAQ on the landing page should answer the inevitable question head-on."
)

add_quote('"Will there be a community / chat?"')

add_para("Not for now. Four reasons:")

reasons = [
    ("Noise reduction.", "No herd mentality, no group-think distorting your read on the market."),
    ("No paid moderators.", "AI moderation is still imperfect, and we refuse to charge users for a chat we cannot reliably regulate."),
    ("Emotional contagion.", "Other people's panic and euphoria are the single biggest leak in most traders' P&L. We will not pipe that into your screen."),
    ("Privacy & intellectual property.", "Trade discussions belong in closed 1v1 consultation, not a public room where your edge becomes everyone's edge."),
]
for i, (title, desc) in enumerate(reasons, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. {title} ")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.color.rgb = ACCENT_DARK
    run = p.add_run(desc)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_TEXT
    p.paragraph_format.space_after = Pt(8)

# ═══════════════════════════════════════════════════════════
# 9. LIABILITY AGREEMENT
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled("9. Liability Agreement", level=1)

add_para(
    "At the end of the signup / upgrade flow, every user must sign a liability agreement that "
    "explicitly releases ANTS from responsibility for any trading losses incurred from acting on "
    "platform content."
)

add_info_box(
    "This is non-optional. No agreement, no access. The liability waiver is the final step before "
    "paid features unlock and is versioned — users re-sign on material changes.",
    accent=True,
)

# ═══════════════════════════════════════════════════════════
# 10. KEY SURFACES & SCREENS
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled("10. Key Surfaces & Screens", level=1)

add_styled_table(
    ["#", "Surface", "Where", "Description"],
    [
        ["1", "Landing Page", "Web", "Positioning, who-is/who-isn't, FAQ, signup CTA, route to Telegram"],
        ["2", "Main Channel", "Telegram", "News + impact analysis, calendar, my channel, psychology primers"],
        ["3", "Stream Mirror", "Telegram", "Live streams (if not hosted on web), conversion surface"],
        ["4", "Collab Channels", "Telegram", "Dedicated influencer / partner broadcast channels"],
        ["5", "DOMAIN Dashboard", "Web", "Paid-tier home — Trading Plan, Consultation entry, Education, hashtag explorer"],
        ["6", "Trading Plan View", "Web", "Gated. Daily / weekly directional plan with setups and levels"],
        ["7", "1v1 Consultation", "Web", "Private chat — AI + team — with hashtag-tagged session history"],
        ["8", "Education Library", "Web", "Structured trading curriculum + psychology primers, partially gated"],
        ["9", "Hashtag Explorer", "Web", "Cross-cut view: every primer, post, and consultation log under a tag"],
        ["10", "Liability Agreement", "Web", "Forced step at end of signup / upgrade"],
        ["11", "Subscription Management", "Web", "Tier status, 3-month renewal, payment history"],
        ["12", "Auth Pages", "Web", "Sign up, login, password reset"],
    ],
    col_widths=[0.4, 1.7, 0.9, 3.5],
)

# ═══════════════════════════════════════════════════════════
# 11. USER ROLES
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled("11. User Roles", level=1)

add_styled_table(
    ["Role", "Capabilities"],
    [
        ["Visitor", "Landing page + free Telegram channels"],
        ["Free User", "Above + free education content + hashtag explorer (limited)"],
        ["Paid Trader (e.g. VIP+ Trader)", "Above + Trading Plan + 1v1 Consultation + full Education + all collab channels"],
        ["Team / Author", "Publish to channels and DOMAIN, run consultation sessions, tag content"],
        ["Influencer / Collaborator", "Publish to a single dedicated channel, scoped to their own surface"],
        ["Admin", "Tier pricing, channel provisioning, liability agreement versions, user / role management"],
    ],
    col_widths=[2.2, 4.3],
)

# ═══════════════════════════════════════════════════════════
# 12. OPEN QUESTIONS
# ═══════════════════════════════════════════════════════════

add_heading_styled("12. Open Questions", level=1)

questions = [
    'Exact tier names and 3-month pricing (what does "VIP+ Trader" actually cost?)',
    "Which features sit at which tier (single paid tier vs. ladder?)",
    "AI provider and model for consultation (and cost ceiling per user)",
    "Where streaming is hosted at launch — web or Telegram-only?",
    "Liability agreement — drafted in-house or run past a lawyer first?",
    "Education library scope at launch (how many primers / lessons day-one?)",
    "Initial hashtag taxonomy — which tags ship in v1?",
    "Collab channel onboarding — how is a new influencer vetted and provisioned?",
    "Domain name (ants.* / domain.* / alpha-network.* — TBD)",
    "Free-tier ceiling — how much do we give away before the paywall hits?",
]

for q in questions:
    p = doc.add_paragraph()
    run = p.add_run(f"[ ]  {q}")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT
    p.paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="8" w:space="4" w:color="E2E8F0"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)
run = p.add_run("ANTS — Alpha Network Trading System | DOMAIN | Product Specification v2.0 | April 2026")
run.font.size = Pt(9)
run.font.color.rgb = MED_GRAY
run.font.name = 'Calibri'
run.font.italic = True

# ── Save ──
output_path = os.path.join(SCRIPT_DIR, "ANTS_DOMAIN_Product_Document.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
